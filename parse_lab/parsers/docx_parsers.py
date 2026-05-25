"""DOCX parser implementations."""

from pathlib import Path

from parse_lab.base import BaseDocumentParser
from parse_lab.parsers._utils import headings_from_markdown, run_timed
from parse_lab.types import ParsedHeading, ParsedTable


class Docx2txtParser(BaseDocumentParser):
    name = "docx2txt"
    description = "docx2txt 基线（langchain Docx2txtLoader）"
    file_types = (".docx",)

    def parse(self, file_path: Path):
        def _do():
            from langchain_community.document_loaders import Docx2txtLoader

            docs = Docx2txtLoader(str(file_path)).load()
            text = "\n\n".join(d.page_content for d in docs)
            return text, [], [], {}

        return run_timed(self.name, file_path, _do)


class PythonDocxParser(BaseDocumentParser):
    name = "python_docx"
    description = "python-docx 解析段落、标题级别与表格"
    file_types = (".docx",)

    def parse(self, file_path: Path):
        def _do():
            from docx import Document as DocxDocument

            doc = DocxDocument(str(file_path))
            parts: list[str] = []
            headings: list[ParsedHeading] = []
            tables: list[ParsedTable] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = (para.style.name or "").lower() if para.style else ""
                level = 0
                if "heading" in style_name:
                    for n in range(1, 7):
                        if str(n) in style_name:
                            level = n
                            break
                    if level:
                        headings.append(
                            ParsedHeading(level=level, text=text, source="docx_style")
                        )
                parts.append(text)

            for tbl in doc.tables:
                rows: list[list[str]] = []
                for row in tbl.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                if rows:
                    tables.append(ParsedTable(rows=rows, source="docx"))

            text = "\n\n".join(parts)
            return text, tables, headings, {"paragraphs": len(parts)}

        return run_timed(self.name, file_path, _do)


class MammothParser(BaseDocumentParser):
    name = "mammoth"
    description = "mammoth 转 Markdown/HTML，正文样式处理较好"
    file_types = (".docx",)
    optional = True

    def parse(self, file_path: Path):
        def _do():
            import mammoth

            with open(file_path, "rb") as f:
                result = mammoth.convert_to_markdown(f)
            text = result.value
            headings = headings_from_markdown(text)
            warnings = [str(w) for w in (result.messages or [])[:5]]
            return text, [], headings, {"warnings": warnings, "format": "markdown"}

        return run_timed(self.name, file_path, _do)
